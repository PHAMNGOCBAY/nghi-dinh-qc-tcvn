import os
import json
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT

JSON_FILE = r"G:\My Drive\NGHI DINH-QC-TCVN\so sánh EIR và 347-348-HTM.json"
OUTPUT_DOCX = r"G:\My Drive\NGHI DINH-QC-TCVN\so sánh EIR và 347-348-HTM.docx"

def create_docx():
    if not os.path.exists(JSON_FILE):
        print(f"Error: File not found {JSON_FILE}")
        return

    with open(JSON_FILE, 'r', encoding='utf-8') as f:
        data = json.load(f)

    doc = Document()
    
    # Set page orientation to landscape
    section = doc.sections[0]
    new_width, new_height = section.page_height, section.page_width
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = new_width
    section.page_height = new_height
    
    # Thêm tiêu đề
    title = doc.add_heading('Báo Cáo So Sánh EIR và Quyết định 347, 348 (Dự án HTM)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Báo cáo được sinh tự động dựa trên từ khóa và phân tích cấu trúc tài liệu.')
    
    # Tạo bảng
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    
    # Header
    hdr_cells = table.rows[0].cells
    headers = ["Tiêu chí", "EIR Dự án HTM", "QĐ 348 (Hướng dẫn chung)", "QĐ 347 (Hướng dẫn chi tiết)", "Đánh giá sơ bộ"]
    for i, header_text in enumerate(headers):
        hdr_cells[i].text = header_text
        # In đậm header
        for run in hdr_cells[i].paragraphs[0].runs:
            run.font.bold = True
            
    # Thêm dữ liệu
    for row_data in data:
        row_cells = table.add_row().cells
        row_cells[0].text = row_data.get("Tiêu chí", "")
        row_cells[1].text = row_data.get("EIR Dự án HTM", "").replace("**", "")
        row_cells[2].text = row_data.get("QĐ 348 (Hướng dẫn chung)", "").replace("**", "")
        row_cells[3].text = row_data.get("QĐ 347 (Hướng dẫn chi tiết)", "").replace("**", "")
        row_cells[4].text = row_data.get("Đánh giá sơ bộ", "")
        
    doc.save(OUTPUT_DOCX)
    print(f"Saved DOCX: {OUTPUT_DOCX}")

if __name__ == "__main__":
    create_docx()
